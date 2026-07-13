from __future__ import annotations

from pathlib import Path

from .base import BaseDocumentLoader


class DocxLoader(BaseDocumentLoader):
    supported_extensions = ["docx"]
    file_type_label = "docx"

    def _extract_text(self, path: Path) -> tuple[str, list[str]]:
        try:
            from docx import Document
        except ImportError as e:  # pragma: no cover
            from ....core.exceptions import DocumentParseError
            raise DocumentParseError("请先安装 python-docx") from e

        doc = Document(str(path))
        lines: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text.strip())
        # 表格
        for table in doc.tables:
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_texts:
                    lines.append(" | ".join(row_texts))
        text = "\n".join(lines)
        pages = [text]  # docx 无原生分页
        return text, pages
